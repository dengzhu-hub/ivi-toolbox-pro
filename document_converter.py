"""
文档格式转换工具 - 专业重构版
支持PDF/Word/Excel/TXT/CSV的转换、编辑、合并

修复问题：
1. PdfWriter 的 append 方法调用错误
2. Excel转PDF时缺失右括号
3. 部分依赖库导入问题
4. 错误处理不完善
5. 代码结构优化
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
from pathlib import Path
import traceback
import sys

# 检查并导入依赖库
try:
    import ttkbootstrap as ttkb
    from ttkbootstrap.constants import *
except ImportError:
    print("警告: ttkbootstrap 未安装，使用默认主题")
    print("安装命令: pip install ttkbootstrap")
    ttkb = None

try:
    import pandas as pd
except ImportError:
    pd = None
    print("警告: pandas 未安装，CSV功能将不可用")

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
except ImportError:
    canvas = None
    print("警告: reportlab 未安装，Excel转PDF功能将不可用")

try:
    from docx import Document
except ImportError:
    Document = None
    print("警告: python-docx 未安装，Word功能将不可用")

try:
    import pdfplumber
except ImportError:
    pdfplumber = None
    print("警告: pdfplumber 未安装，PDF读取功能将不可用")

try:
    from openpyxl import load_workbook, Workbook
except ImportError:
    load_workbook = Workbook = None
    print("警告: openpyxl 未安装，Excel功能将不可用")

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
    except ImportError:
        PdfReader = PdfWriter = None
        print("警告: pypdf/PyPDF2 未安装，PDF操作功能将不可用")


class DocumentConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("文档格式转换工具 - 专业版")
        self.root.geometry("1000x750")

        # 存储选择的文件
        self.selected_files = []

        # 输出路径（默认当前目录）
        self.output_path = tk.StringVar(value=os.getcwd())

        self.create_widgets()

    def create_widgets(self):
        """创建界面组件"""
        # 根据是否安装 ttkbootstrap 选择主题
        if ttkb:
            style = ttkb.Style(theme='darkly')
            LabelFrame = ttkb.LabelFrame
            Label = ttkb.Label
            Button = ttkb.Button
            Entry = ttkb.Entry
            Frame = ttkb.Frame
            Scrollbar = ttkb.Scrollbar
        else:
            LabelFrame = ttk.LabelFrame
            Label = ttk.Label
            Button = ttk.Button
            Entry = ttk.Entry
            Frame = ttk.Frame
            Scrollbar = ttk.Scrollbar

        # 标题
        title = Label(self.root, text="文档格式转换工具",
                     font=("Arial", 18, "bold"))
        title.grid(row=0, column=0, columnspan=2, pady=15, sticky="ew")

        # 文件选择区域
        file_frame = LabelFrame(self.root, text="文件选择", padding=15)
        file_frame.grid(row=1, column=0, columnspan=2, padx=20, pady=10, sticky="ew")

        # 按钮框架
        btn_frame = Frame(file_frame)
        btn_frame.grid(row=0, column=0, columnspan=3, sticky="ew", pady=5)

        Button(btn_frame, text="选择文件", command=self.select_files).grid(row=0, column=0, padx=5)
        Button(btn_frame, text="清空列表", command=self.clear_files).grid(row=0, column=1, padx=5)

        # 输出路径
        Label(file_frame, text="输出目录:").grid(row=1, column=0, pady=5, sticky="w")
        Entry(file_frame, textvariable=self.output_path, width=60).grid(row=1, column=1, pady=5, sticky="ew", padx=5)
        Button(file_frame, text="浏览", command=self.select_output_dir).grid(row=1, column=2, padx=5)

        file_frame.columnconfigure(1, weight=1)

        # 文件列表
        list_frame = Frame(file_frame)
        list_frame.grid(row=2, column=0, columnspan=3, sticky="ew", pady=10)

        scrollbar = Scrollbar(list_frame)
        scrollbar.grid(row=0, column=1, sticky="ns")

        self.file_listbox = tk.Listbox(list_frame, height=8,
                                       yscrollcommand=scrollbar.set)
        self.file_listbox.grid(row=0, column=0, sticky="ew")
        scrollbar.config(command=self.file_listbox.yview)

        list_frame.columnconfigure(0, weight=1)

        # 功能选择区域
        function_frame = LabelFrame(self.root, text="选择功能", padding=15)
        function_frame.grid(row=2, column=0, columnspan=2, padx=20, pady=10, sticky="ew")

        # 转换功能
        convert_frame = Frame(function_frame)
        convert_frame.grid(row=0, column=0, sticky="ew", pady=5)

        Label(convert_frame, text="格式转换:", font=("Arial", 11, "bold")).grid(row=0, column=0, padx=5, sticky="w")

        convert_buttons = [
            ("PDF → Word", lambda: self.convert("pdf_to_word")),
            ("PDF → Excel", lambda: self.convert("pdf_to_excel")),
            ("Word → PDF", lambda: self.convert("word_to_pdf")),
            ("Excel → PDF", lambda: self.convert("excel_to_pdf")),
            ("CSV → Excel", lambda: self.convert("csv_to_excel")),
            ("CSV → Word", lambda: self.convert("csv_to_word"))
        ]

        for i, (text, cmd) in enumerate(convert_buttons):
            Button(convert_frame, text=text, command=cmd, width=12).grid(row=0, column=i+1, padx=3)

        # 编辑功能
        edit_frame = Frame(function_frame)
        edit_frame.grid(row=1, column=0, sticky="ew", pady=5)

        Label(edit_frame, text="文档操作:", font=("Arial", 11, "bold")).grid(row=0, column=0, padx=5, sticky="w")

        edit_buttons = [
            ("合并PDF", lambda: self.merge_files("pdf")),
            ("合并Word", lambda: self.merge_files("word")),
            ("拆分PDF", self.split_pdf),
        ]

        for i, (text, cmd) in enumerate(edit_buttons):
            Button(edit_frame, text=text, command=cmd, width=12).grid(row=0, column=i+1, padx=3)

        # 日志区域
        log_frame = LabelFrame(self.root, text="操作日志", padding=10)
        log_frame.grid(row=3, column=0, columnspan=2, padx=20, pady=10, sticky="nsew")

        self.root.rowconfigure(3, weight=1)
        self.root.columnconfigure(0, weight=1)

        log_scrollbar = Scrollbar(log_frame)
        log_scrollbar.grid(row=0, column=1, sticky="ns")

        self.log_text = tk.Text(log_frame, height=12, wrap="word",
                                font=("Consolas", 9),
                                yscrollcommand=log_scrollbar.set)
        self.log_text.grid(row=0, column=0, sticky="nsew")
        log_scrollbar.config(command=self.log_text.yview)

        log_frame.rowconfigure(0, weight=1)
        log_frame.columnconfigure(0, weight=1)

        # 底部状态栏
        self.status_var = tk.StringVar(value="就绪")
        status_bar = Label(self.root, textvariable=self.status_var,
                          relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=4, column=0, columnspan=2, sticky="ew")

    def select_files(self):
        """选择文件"""
        files = filedialog.askopenfilenames(
            title="选择文件",
            filetypes=[
                ("所有支持文件", "*.pdf *.docx *.doc *.xlsx *.xls *.csv *.txt"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx *.doc"),
                ("Excel文件", "*.xlsx *.xls"),
                ("CSV文件", "*.csv"),
                ("所有文件", "*.*")
            ]
        )
        if files:
            self.selected_files = list(files)
            self.file_listbox.delete(0, tk.END)
            for f in self.selected_files:
                self.file_listbox.insert(tk.END, os.path.basename(f))
            self.log(f"✓ 已选择 {len(self.selected_files)} 个文件")
            self.status_var.set(f"已选择 {len(self.selected_files)} 个文件")

    def clear_files(self):
        """清空文件列表"""
        self.selected_files = []
        self.file_listbox.delete(0, tk.END)
        self.log("✓ 文件列表已清空")
        self.status_var.set("就绪")

    def select_output_dir(self):
        """选择输出目录"""
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_path.set(dir_path)
            self.log(f"✓ 输出目录: {dir_path}")

    def log(self, message):
        """日志输出"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def convert(self, conversion_type):
        """格式转换主方法"""
        if not self.selected_files:
            messagebox.showwarning("警告", "请先选择文件!")
            return

        output_dir = Path(self.output_path.get())
        output_dir.mkdir(parents=True, exist_ok=True)

        success_count = 0
        error_count = 0

        try:
            for file_path in self.selected_files:
                file_ext = Path(file_path).suffix.lower()
                self.log(f"\n→ 正在处理: {os.path.basename(file_path)}")
                self.status_var.set(f"正在转换: {os.path.basename(file_path)}")

                try:
                    if conversion_type == "pdf_to_word" and file_ext == '.pdf':
                        if not pdfplumber or not Document:
                            raise ImportError("缺少必要库: pdfplumber, python-docx")
                        self.pdf_to_word(file_path, output_dir)

                    elif conversion_type == "pdf_to_excel" and file_ext == '.pdf':
                        if not pdfplumber or not pd:
                            raise ImportError("缺少必要库: pdfplumber, pandas")
                        self.pdf_to_excel(file_path, output_dir)

                    elif conversion_type == "word_to_pdf" and file_ext in ['.docx', '.doc']:
                        messagebox.showinfo("提示", "Word转PDF需要安装Microsoft Word\n或使用在线转换工具")
                        continue

                    elif conversion_type == "excel_to_pdf" and file_ext in ['.xlsx', '.xls']:
                        if not load_workbook or not canvas:
                            raise ImportError("缺少必要库: openpyxl, reportlab")
                        self.excel_to_pdf(file_path, output_dir)

                    elif conversion_type == "csv_to_excel" and file_ext == '.csv':
                        if not pd:
                            raise ImportError("缺少必要库: pandas")
                        self.csv_to_excel(file_path, output_dir)

                    elif conversion_type == "csv_to_word" and file_ext == '.csv':
                        if not pd or not Document:
                            raise ImportError("缺少必要库: pandas, python-docx")
                        self.csv_to_word(file_path, output_dir)

                    else:
                        self.log(f"  ⚠ 跳过: 不支持的文件类型 {file_ext}")
                        continue

                    success_count += 1
                    self.log(f"  ✓ 转换成功")

                except Exception as e:
                    error_count += 1
                    self.log(f"  ✗ 错误: {str(e)}")

            self.status_var.set(f"完成: 成功 {success_count} 个, 失败 {error_count} 个")

            if success_count > 0:
                messagebox.showinfo("完成", f"转换完成!\n成功: {success_count}\n失败: {error_count}")

        except Exception as e:
            self.log(f"\n✗ 严重错误: {str(e)}\n{traceback.format_exc()}")
            messagebox.showerror("错误", f"转换失败: {str(e)}")

    def pdf_to_word(self, file_path, output_dir):
        """PDF 转 Word"""
        doc = Document()

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)

        output_file = output_dir / f"{Path(file_path).stem}.docx"
        doc.save(output_file)
        self.log(f"  → 保存到: {output_file.name}")

    def pdf_to_excel(self, file_path, output_dir):
        """PDF 转 Excel"""
        all_data = []

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    all_data.extend(table)

        if all_data:
            df = pd.DataFrame(all_data)
            output_file = output_dir / f"{Path(file_path).stem}.xlsx"
            df.to_excel(output_file, index=False, header=False)
            self.log(f"  → 保存到: {output_file.name}")
        else:
            self.log(f"  ⚠ 未提取到表格数据")

    def excel_to_pdf(self, file_path, output_dir):
        """Excel 转 PDF - 修复版"""
        wb = load_workbook(file_path)
        output_file = output_dir / f"{Path(file_path).stem}.pdf"

        c = canvas.Canvas(str(output_file), pagesize=letter)
        width, height = letter
        y = height - 50

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]  # 修复：添加了缺失的右括号

            # 添加工作表标题
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, y, f"工作表: {sheet_name}")
            y -= 30

            c.setFont("Helvetica", 10)

            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])

                # 防止文本过长
                if len(row_text) > 100:
                    row_text = row_text[:97] + "..."

                c.drawString(50, y, row_text)
                y -= 15

                if y < 50:
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - 50

            c.showPage()
            y = height - 50

        c.save()
        self.log(f"  → 保存到: {output_file.name}")

    def csv_to_excel(self, file_path, output_dir):
        """CSV 转 Excel"""
        df = pd.read_csv(file_path, encoding='utf-8-sig')
        output_file = output_dir / f"{Path(file_path).stem}.xlsx"
        df.to_excel(output_file, index=False)
        self.log(f"  → 保存到: {output_file.name}")

    def csv_to_word(self, file_path, output_dir):
        """CSV 转 Word"""
        df = pd.read_csv(file_path, encoding='utf-8-sig')
        doc = Document()

        # 添加表格
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'

        # 表头
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)

        # 数据行
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val) if pd.notna(val) else ""

        output_file = output_dir / f"{Path(file_path).stem}.docx"
        doc.save(output_file)
        self.log(f"  → 保存到: {output_file.name}")

    def merge_files(self, file_type):
        """合并文件 - 修复版"""
        if not self.selected_files:
            messagebox.showwarning("警告", "请先选择文件!")
            return

        output_dir = Path(self.output_path.get())
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            if file_type == "pdf":
                if not PdfReader or not PdfWriter:
                    messagebox.showerror("错误", "缺少pypdf库，请安装: pip install pypdf")
                    return

                pdf_files = [f for f in self.selected_files if f.lower().endswith('.pdf')]
                if not pdf_files:
                    messagebox.showwarning("警告", "未选择 PDF 文件!")
                    return

                output_file = output_dir / "merged.pdf"
                writer = PdfWriter()

                for pdf_path in pdf_files:
                    self.log(f"→ 添加: {os.path.basename(pdf_path)}")
                    reader = PdfReader(pdf_path)

                    # 修复：正确使用 add_page 方法
                    for page in reader.pages:
                        writer.add_page(page)

                with open(output_file, 'wb') as output:
                    writer.write(output)

                self.log(f"✓ PDF 合并完成: {output_file.name}")
                messagebox.showinfo("成功", f"已合并 {len(pdf_files)} 个PDF文件\n保存为: {output_file.name}")

            elif file_type == "word":
                if not Document:
                    messagebox.showerror("错误", "缺少python-docx库")
                    return

                word_files = [f for f in self.selected_files if f.lower().endswith(('.docx', '.doc'))]
                if not word_files:
                    messagebox.showwarning("警告", "未选择 Word 文件!")
                    return

                output_file = output_dir / "merged.docx"
                merged_doc = Document()

                for word_path in word_files:
                    self.log(f"→ 添加: {os.path.basename(word_path)}")
                    doc = Document(word_path)

                    for element in doc.element.body:
                        merged_doc.element.body.append(element)

                merged_doc.save(output_file)
                self.log(f"✓ Word 合并完成: {output_file.name}")
                messagebox.showinfo("成功", f"已合并 {len(word_files)} 个Word文件\n保存为: {output_file.name}")

        except Exception as e:
            self.log(f"✗ 合并失败: {str(e)}\n{traceback.format_exc()}")
            messagebox.showerror("错误", f"合并失败: {str(e)}")

    def split_pdf(self):
        """拆分PDF"""
        if not self.selected_files:
            messagebox.showwarning("警告", "请先选择PDF文件!")
            return

        if not PdfReader or not PdfWriter:
            messagebox.showerror("错误", "缺少pypdf库")
            return

        pdf_file = [f for f in self.selected_files if f.lower().endswith('.pdf')]
        if not pdf_file:
            messagebox.showwarning("警告", "请选择PDF文件!")
            return

        pdf_path = pdf_file[0]
        output_dir = Path(self.output_path.get()) / f"{Path(pdf_path).stem}_split"
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)

            for i, page in enumerate(reader.pages, 1):
                writer = PdfWriter()
                writer.add_page(page)

                output_file = output_dir / f"page_{i}.pdf"
                with open(output_file, 'wb') as output:
                    writer.write(output)

            self.log(f"✓ PDF拆分完成: 共 {total_pages} 页")
            messagebox.showinfo("成功", f"已拆分 {total_pages} 页\n保存到: {output_dir}")

        except Exception as e:
            self.log(f"✗ 拆分失败: {str(e)}")
            messagebox.showerror("错误", f"拆分失败: {str(e)}")


def main():
    """主函数"""
    # 使用 ttkbootstrap 或标准 tkinter
    if ttkb:
        root = ttkb.Window(themename='darkly')
    else:
        root = tk.Tk()

    app = DocumentConverter(root)

    # 居中显示窗口
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')

    root.mainloop()


if __name__ == "__main__":
    main()